# report_gen.py — python-docx: JSON → filled .docx report
import os, re, copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE_PATH = os.getenv('TEMPLATE_PATH', 'casad_template.docx')
OUTPUT_DIR    = os.getenv('OUTPUT_DIR', 'media')


def _fill_placeholders(doc: Document, data: dict) -> None:
    """Replace {{key}} placeholders in all paragraphs and table cells."""
    flat = {}
    for k, v in data.items():
        if isinstance(v, list):
            flat[k] = '\n'.join(str(i) for i in v)
        elif isinstance(v, dict):
            for sk, sv in v.items():
                flat[f'{k}.{sk}'] = str(sv) if sv else ''
        else:
            flat[k] = str(v) if v else ''

    def replace_in_para(para):
        full_text = ''.join(r.text for r in para.runs)
        if '{{' not in full_text:
            return
        new_text = full_text
        for key, val in flat.items():
            new_text = new_text.replace(f'{{{{{key}}}}}', val)
        if new_text != full_text and para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ''

    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)


def _add_bookmark(para, bookmark_id: int, name: str) -> None:
    """Wrap paragraph content in a named bookmark."""
    p = para._p
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))
    p.insert(0, bm_start)
    p.append(bm_end)


def _replace_photo_refs_with_hyperlinks(doc: Document) -> None:
    """Replace '(Photo No.-X)' text in table cells with internal bookmark hyperlinks."""
    import re
    pattern = re.compile(r'(\(Photo No\.-\d+\))')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = ''.join(r.text for r in para.runs)
                    if '(Photo No.-' not in full_text:
                        continue

                    # Capture run formatting from first run
                    p = para._p
                    first_rpr = None
                    first_r = p.find(qn('w:r'))
                    if first_r is not None:
                        first_rpr = first_r.find(qn('w:rPr'))

                    # Remove all existing runs
                    for r_el in p.findall(qn('w:r')):
                        p.remove(r_el)

                    # Rebuild paragraph — plain text + hyperlink elements
                    for part in pattern.split(full_text):
                        if not part:
                            continue
                        m = re.match(r'\(Photo No\.-(\d+)\)', part)
                        if m:
                            fig_num  = m.group(1)
                            anchor   = f'figure_{fig_num}'
                            hl = OxmlElement('w:hyperlink')
                            hl.set(qn('w:anchor'), anchor)
                            r_el = OxmlElement('w:r')
                            # Explicit hyperlink formatting: blue + underline
                            rPr  = OxmlElement('w:rPr')
                            color_el = OxmlElement('w:color')
                            color_el.set(qn('w:val'), '0563C1')
                            rPr.append(color_el)
                            u_el = OxmlElement('w:u')
                            u_el.set(qn('w:val'), 'single')
                            rPr.append(u_el)
                            r_el.append(rPr)
                            t_el = OxmlElement('w:t')
                            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                            t_el.text = part
                            r_el.append(t_el)
                            hl.append(r_el)
                            p.append(hl)
                        else:
                            r_el = OxmlElement('w:r')
                            if first_rpr is not None:
                                import copy
                                r_el.append(copy.deepcopy(first_rpr))
                            t_el = OxmlElement('w:t')
                            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                            t_el.text = part
                            r_el.append(t_el)
                            p.append(r_el)


def _shade_cell(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _add_defect_circle(para, x_pct: float, y_pct: float,
                        img_w_emu: int, img_h_emu: int, shape_id: int) -> None:
    """Append an editable red oval DrawingML shape (wp:anchor) to the paragraph.

    The shape floats over the image at the detected defect location.
    Engineers can select, move, or resize it in Word like any shape.
    """
    from lxml import etree

    RADIUS = max(400000, min(img_w_emu, img_h_emu) // 7)  # ~proportional to image size
    space_before_emu = 76200  # 6pt space_before on img_para

    cx = int(x_pct * img_w_emu)
    cy = space_before_emu + int(y_pct * img_h_emu)

    left = max(0, cx - RADIUS)
    top  = max(0, cy - RADIUS)
    diam = RADIUS * 2

    xml = (
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        f' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        f' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f'<w:drawing>'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        f' simplePos="0" relativeHeight="251658240" behindDoc="0"'
        f' locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="column"><wp:posOffset>{left}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>{top}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{diam}" cy="{diam}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{shape_id}" name="DefectCircle{shape_id}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic>'
        f'<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f'<wps:wsp>'
        f'<wps:cNvSpPr><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>'
        f'<wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{diam}" cy="{diam}"/></a:xfrm>'
        f'<a:prstGeom prst="ellipse"><a:avLst/></a:prstGeom>'
        f'<a:noFill/>'
        f'<a:ln w="57150" cmpd="sng"><a:solidFill><a:srgbClr val="FF0000"/></a:solidFill></a:ln>'
        f'</wps:spPr>'
        f'<wps:bodyPr/>'
        f'</wps:wsp>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:anchor>'
        f'</w:drawing>'
        f'</w:r>'
    )
    para._p.append(etree.fromstring(xml))


def _insert_photos_at_marker(doc: Document, marker: str, photos: list,
                              titles: list = None, coords: list = None,
                              fig_offset: int = 0,
                              show_figure_label: bool = True) -> None:
    """Find the marker paragraph and replace it with bordered photo tables."""
    target = None
    for para in doc.paragraphs:
        if para.text.strip() == marker:
            target = para
            break
    if target is None:
        return

    parent = target._element.getparent()
    insert_pos = list(parent).index(target._element)
    parent.remove(target._element)

    valid = [(p,
              titles[i] if titles and i < len(titles) else '',
              coords[i] if coords and i < len(coords) else None)
             for i, p in enumerate(photos)
             if p and os.path.exists(p)]

    shape_id_counter = fig_offset * 100 + 1   # unique IDs across appendices

    for idx, (photo_path, title, photo_coords) in enumerate(valid, 1):
        if show_figure_label:
            cap_text = f'Figure {fig_offset + idx}'
            if title:
                cap_text += f':  {title}'
        else:
            cap_text = title  # title only, no "Figure N:" prefix

        # --- Bordered 2-row table: [image row] / [caption row] ---
        tbl = doc.add_table(rows=2, cols=1)
        tbl.style = 'Table Grid'

        # Row 0 — photo centred with padding
        img_cell = tbl.rows[0].cells[0]
        img_para = img_cell.paragraphs[0]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.space_before = Pt(6)
        img_para.paragraph_format.space_after  = Pt(6)

        # Constrain to fit within page: max 5.5" wide, 4.0" tall
        from PIL import Image as _PIL
        with _PIL.open(photo_path) as _img:
            w_px, h_px = _img.size
        MAX_W = Inches(5.5)
        MAX_H = Inches(4.0)
        if w_px > 0 and h_px / w_px * MAX_W > MAX_H:
            img_para.add_run().add_picture(photo_path, height=MAX_H)
            img_w_emu = int(MAX_H * w_px / h_px)
            img_h_emu = int(MAX_H)
        else:
            img_para.add_run().add_picture(photo_path, width=MAX_W)
            img_w_emu = int(MAX_W)
            img_h_emu = int(MAX_W * h_px / w_px)

        # Overlay an editable red oval at the detected defect location
        if photo_coords:
            try:
                _add_defect_circle(img_para,
                                   photo_coords[0], photo_coords[1],
                                   img_w_emu, img_h_emu,
                                   shape_id=shape_id_counter)
                shape_id_counter += 1
            except Exception as e:
                print(f"ADD CIRCLE SHAPE FAILED for {photo_path}: {e}")

        # Row 1 — caption, grey background
        cap_cell = tbl.rows[1].cells[0]
        _shade_cell(cap_cell, 'EBF5FB')
        cap_para = cap_cell.paragraphs[0]
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_before = Pt(4)
        cap_para.paragraph_format.space_after  = Pt(4)
        cap_run = cap_para.add_run(cap_text)
        cap_run.italic     = True
        cap_run.bold       = True
        cap_run.font.size  = Pt(9)

        # Add bookmark so (Photo No.-X) hyperlinks can navigate here
        if show_figure_label:
            _add_bookmark(cap_para, fig_offset + idx, f'figure_{fig_offset + idx}')

        # Move table to correct position (relationships stay in main doc ✓)
        tbl_elem = tbl._element
        parent.remove(tbl_elem)
        parent.insert(insert_pos, tbl_elem)
        insert_pos += 1

        # Small spacer paragraph between tables
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(4)
        sp_elem = sp._element
        parent.remove(sp_elem)
        parent.insert(insert_pos, sp_elem)
        insert_pos += 1

        # Page break after every 2 photos (except the last)
        if idx % 2 == 0 and idx < len(valid):
            pb = doc.add_paragraph()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            pb.add_run()._r.append(br)
            pb_elem = pb._element
            parent.remove(pb_elem)
            parent.insert(insert_pos, pb_elem)
            insert_pos += 1


def _ensure_fixed_layout(doc: Document) -> None:
    """Force tblLayout=fixed on all tables so Word respects column widths."""
    for table in doc.tables:
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tblPr)
        existing = tblPr.find(qn('w:tblLayout'))
        if existing is not None:
            tblPr.remove(existing)
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)


def build_docx(report_json: dict) -> str:
    """Fill CASAD template with report_json and return saved file path."""
    doc = Document(TEMPLATE_PATH)
    _fill_placeholders(doc, report_json)
    _ensure_fixed_layout(doc)

    raw_photos     = report_json.get('photos', [])
    raw_titles     = report_json.get('photo_titles', [])
    raw_categories = report_json.get('photo_categories', [])
    raw_coords     = report_json.get('photo_coords', [])

    # Safety: restore any missing photo files from session BLOB data
    session_messages = report_json.get('_messages', [])
    for m in session_messages:
        path = m.get('media_path')
        blob = m.get('image_data')
        if path and blob and not os.path.exists(path):
            try:
                os.makedirs(os.path.dirname(path), exist_ok=True)
                with open(path, 'wb') as f:
                    f.write(blob)
                print(f"BUILD_DOCX restored: {path}")
            except Exception as e:
                print(f"BUILD_DOCX restore failed for {path}: {e}")

    # Pad lists to same length
    n = len(raw_photos)
    raw_titles     = list(raw_titles)     + [''] * n
    raw_categories = list(raw_categories) + ['damage'] * n
    raw_coords     = list(raw_coords)     + [None] * n

    # Split into general and damage buckets, keeping titles/coords aligned
    general_photos, general_titles, general_coords = [], [], []
    damage_photos,  damage_titles,  damage_coords  = [], [], []

    for path, title, cat, coord in zip(raw_photos, raw_titles, raw_categories, raw_coords):
        if not path or not os.path.exists(path):
            continue
        if str(cat).lower() in ('damage', 'damaged'):
            damage_photos.append(path)
            damage_titles.append(title)
            damage_coords.append(coord)
        else:
            general_photos.append(path)
            general_titles.append(title)
            general_coords.append(coord)

    print(f"BUILD_DOCX general photos: {general_photos}")
    print(f"BUILD_DOCX damage  photos: {damage_photos}")

    # Insert into Appendix A (general) — no figure numbering, no defect circles
    if general_photos:
        _insert_photos_at_marker(doc, '[[PHOTO_APPENDIX_A]]', general_photos, general_titles,
                                  coords=None, fig_offset=0, show_figure_label=False)
    else:
        for para in doc.paragraphs:
            if para.text.strip() == '[[PHOTO_APPENDIX_A]]':
                para.runs[0].text = 'No general photographs submitted.'
                break

    # Insert into Appendix B (damage) — figure numbers start at 1, with editable circles
    if damage_photos:
        _insert_photos_at_marker(doc, '[[PHOTO_APPENDIX_B]]', damage_photos, damage_titles,
                                  coords=damage_coords, fig_offset=0)
    else:
        for para in doc.paragraphs:
            if para.text.strip() == '[[PHOTO_APPENDIX_B]]':
                para.runs[0].text = 'No damage photographs submitted.'
                break

    # Convert (Photo No.-X) text to clickable hyperlinks now that bookmarks exist
    _replace_photo_refs_with_hyperlinks(doc)

    river    = re.sub(r'[^\w\-]', '_', report_json.get('river_name', 'bridge'))
    road     = re.sub(r'[^\w\-]', '_', report_json.get('road_name',  'road'))
    date     = report_json.get('date_of_survey', 'report').replace('/', '-')
    out_path = os.path.join(OUTPUT_DIR, f'CASAD_{river}_{road}_{date}.docx')
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(out_path)
    return out_path
