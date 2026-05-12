"""
Generate casad_template.docx — the CASAD bridge inspection report template.
Run once: python3 create_template.py
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── colour palette ────────────────────────────────────────────────────────────
HEADER_BG   = "1F3864"   # dark navy
SECTION_BG  = "D6E4F0"   # light blue
SUB_BG      = "EBF5FB"   # very light blue
WHITE       = "FFFFFF"
DARK_TEXT   = "1F3864"


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="AAAAAA"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_para(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    for p in cell.paragraphs:
        cell._element.remove(p._element)
    p   = cell.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p


def add_row(table, sr, description, placeholder):
    row   = table.add_row()
    cells = row.cells

    cell_para(cells[0], sr,          size=9, bold=bool(sr))
    cell_para(cells[1], description, size=9)
    cell_para(cells[2], placeholder, size=9)

    for c in cells:
        set_cell_bg(c, WHITE)
        set_cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return row


def section_header(table, title):
    row   = table.add_row()
    cells = row.cells
    cells[0].merge(cells[2])
    cell_para(cells[0], title, bold=True, size=10, color=WHITE, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_bg(cells[0], HEADER_BG)
    set_cell_borders(cells[0], "FFFFFF")


def span_row(table, title, bg=None, bold=True, size=9, color=None):
    """Full-width row spanning all 3 columns — used for sub-section headers."""
    row   = table.add_row()
    cells = row.cells
    cells[0].merge(cells[2])
    kw = dict(bold=bold, size=size)
    if color:
        kw["color"] = color
    cell_para(cells[0], title, **kw)
    set_cell_bg(cells[0], bg or WHITE)
    set_cell_borders(cells[0])


def b_section_header(table, sr, title):
    """Numbered section header row (e.g. '2  SUPERSTRUCTURE') with light blue bg."""
    row   = table.add_row()
    cells = row.cells
    # Merge description + details columns so title spans them
    cells[1].merge(cells[2])
    cell_para(cells[0], sr,    bold=True, size=9, color=DARK_TEXT)
    cell_para(cells[1], title, bold=True, size=9, color=DARK_TEXT)
    for c in (cells[0], cells[1]):
        set_cell_bg(c, SECTION_BG)
        set_cell_borders(c)


# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Page setup — A4, 2 cm margins
sec = doc.sections[0]
sec.page_width  = Cm(21)
sec.page_height = Cm(29.7)
for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(sec, attr, Cm(2))

# ── Title ─────────────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("REPORT FOR BRIDGE SURVEY AND INSPECTION")
r.bold      = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor.from_string(DARK_TEXT)
t.paragraph_format.space_after = Pt(6)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("CASAD CONSULTANTS PVT. LTD.")
r2.bold      = True
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor.from_string(DARK_TEXT)
sub.paragraph_format.space_after = Pt(10)

# ── Main 3-column table ───────────────────────────────────────────────────────
# Column widths (A4 content ~17 cm): SR 0.7 | Description 7.5 | Details 8.8
table = doc.add_table(rows=0, cols=3)
table.style = "Table Grid"

col_widths = [Cm(0.3), Cm(7.9), Cm(8.8)]
for i, w in enumerate(col_widths):
    for cell in table.columns[i].cells:
        cell.width = w

# Column headers
hrow = table.add_row()
for i, (txt, w) in enumerate(zip(["SR\nNO.", "DESCRIPTION", "DETAILS"], col_widths)):
    cell_para(hrow.cells[i], txt, bold=True, size=10, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(hrow.cells[i], HEADER_BG)
    set_cell_borders(hrow.cells[i], "FFFFFF")
    hrow.cells[i].width = w

# ═══ SECTION A ═══════════════════════════════════════════════════════════════
section_header(table, "(A)  DETAILS OF EXISTING BRIDGE")

rows_a = [
    ("1",  "NAME OF RIVER:",                              "{{river_name}}"),
    ("2",  "NAME OF THE ROAD:",                           "{{road_name}}"),
    ("3",  "CHAINAGE OF THE BRIDGE:\nLatitude\nLongitude","{{chainage}}\n{{latitude}}\n{{longitude}}"),
    ("4",  "NAME OF CIRCLE:\nNAME OF DIVISION:\nNAME OF SUB DIVISION:",
                                                           "{{circle}}\n{{division}}\n{{sub_division}}"),
    ("5",  "NO. OF SPAN:",                                "{{no_of_spans}}"),
    ("6",  "SPAN LENGTH & ARRANGEMENT:",                  "{{span_length}}"),
    ("7",  "TYPE OF BRIDGE:\n(Simply Supported / Continuous / Balanced Cantilever / Arch / Other)",
                                                           "{{bridge_type}}"),
    ("8",  "TYPE OF SUPERSTRUCTURE:",                     "{{superstructure_type}}"),
    ("9",  "TYPE OF SUBSTRUCTURE:",                       "{{substructure_type}}"),
    ("10", "TYPE OF FOUNDATION:",                         "{{foundation_type}}"),
    ("11", "TYPE OF BEARING:",                            "{{bearing_type_detail}}"),
    ("12", "TOTAL LENGTH OF BRIDGE:",                     "{{total_length}}"),
    ("13", "TOTAL LENGTH OF APPROACH:",                   "{{approach_length}}"),
    ("14", "TYPE OF RAILING:\n(RCC parapet / Pipe railing / Crash barrier)",
                                                           "{{railing_type}}"),
    ("15", "DETAIL OF RIVER TRAINING WORK IF ANY?",       "{{river_training}}"),
    ("16", "REPAIR / STRENGTHENING WORK DETAIL IF ANY?",  "{{repair_work}}"),
    ("17", "CLEAR CARRIAGEWAY WIDTH:",                    "{{carriageway_width}}"),
    ("18", "YEAR OF CONSTRUCTION:",                       "{{year_of_construction}}"),
    ("19", "SPECIFY HIGH LEVEL / SUBMERSIBLE BRIDGE:",    "{{bridge_level_type}}"),
    ("20", "WHETHER RIVER IS PERENNIAL OR NOT?",          "{{river_perennial}}"),
]
for sr, desc, ph in rows_a:
    add_row(table, sr, desc, ph)

# ═══ SECTION B ═══════════════════════════════════════════════════════════════
section_header(table, "(B)  SURVEY AND CONDITION ASSESSMENT")
add_row(table, "1", "DATE OF SURVEY & INSPECTION:", "{{date_of_survey}}")

# B2 — Superstructure
b_section_header(table, "2", "SUPERSTRUCTURE")
ss_items = [
    ("(a)", "CRACKS (Length, Width, Location):",    "{{ss_cracks}}"),
    ("(b)", "LEACHING:",                             "{{ss_leaching}}"),
    ("(c)", "HONEY COMBING:",                        "{{ss_honey_combing}}"),
    ("(d)", "EXPOSED REINFORCEMENT:",                "{{ss_exposed_rebar}}"),
    ("(e)", "LEAKAGE PATCHES & WATERMARKS:",         "{{ss_leakage_patches}}"),
    ("(f)", "SPALLING:",                             "{{ss_spalling}}"),
    ("(g)", "RUST MARKS:",                           "{{ss_rust_marks}}"),
    ("(h)", "SHUTTERING DEFECTS:",                   "{{ss_shuttering_defects}}"),
    ("(i)", "DELAMINATION:",                         "{{ss_delamination}}"),
    ("(j)", "ANY OTHER SPECIFY:",                    "{{ss_other}}"),
]
for sr, desc, ph in ss_items:
    add_row(table, sr, desc, ph)

# B3 — Substructure
b_section_header(table, "3", "SUBSTRUCTURE")
sub_items = [
    ("(a)", "CRACKS (Length, Width, Location):",    "{{sub_cracks}}"),
    ("(b)", "LEACHING:",                             "{{sub_leaching}}"),
    ("(c)", "HONEY COMBING:",                        "{{sub_honey_combing}}"),
    ("(d)", "EXPOSED REINFORCEMENT:",                "{{sub_exposed_rebar}}"),
    ("(e)", "SPALLING:",                             "{{sub_spalling}}"),
    ("(f)", "RUST MARKS:",                           "{{sub_rust_marks}}"),
    ("(g)", "SHUTTERING DEFECTS:",                   "{{sub_shuttering_defects}}"),
    ("(h)", "DELAMINATION:",                         "{{sub_delamination}}"),
    ("(i)", "TILTING IF ANY:",                       "{{sub_tilting}}"),
    ("(j)", "ANY OTHER SPECIFY:",                    "{{sub_other}}"),
]
for sr, desc, ph in sub_items:
    add_row(table, sr, desc, ph)

# B4 — Foundations
b_section_header(table, "4", "FOUNDATIONS (pile / cap etc.)")
found_items = [
    ("(a)", "CRACKS (Length, Width, Location):",    "{{found_cracks}}"),
    ("(b)", "LEACHING:",                             "{{found_leaching}}"),
    ("(c)", "HONEY COMBING:",                        "{{found_honey_combing}}"),
    ("(d)", "EXPOSED REINFORCEMENT:",                "{{found_exposed_rebar}}"),
    ("(e)", "SPALLING:",                             "{{found_spalling}}"),
    ("(f)", "RUST MARKS:",                           "{{found_rust_marks}}"),
    ("(g)", "SHUTTERING DEFECTS:",                   "{{found_shuttering_defects}}"),
    ("(h)", "DELAMINATION:",                         "{{found_delamination}}"),
    ("(i)", "SETTLEMENT:",                           "{{found_settlement}}"),
    ("(j)", "TILTING IF ANY:",                       "{{found_tilting}}"),
    ("(k)", "SCOUR DETAILS FOR EACH FOUNDATION:",    "{{found_scour}}"),
    ("(l)", "ANY OTHER SPECIFY:",                    "{{found_other}}"),
]
for sr, desc, ph in found_items:
    add_row(table, sr, desc, ph)

# B5 — Bearings
b_section_header(table, "5", "BEARINGS")
for sr, desc, ph in [
    ("(a)", "DISPLACEMENT:",  "{{bearing_displacement}}"),
    ("(b)", "DISTORTION:",    "{{bearing_distortion}}"),
    ("(c)", "CORROSION:",     "{{bearing_corrosion}}"),
]:
    add_row(table, sr, desc, ph)

# B6 — Approach
b_section_header(table, "6", "APPROACH AND OTHER")
for sr, desc, ph in [
    ("(a)", "SETTLEMENT:",         "{{approach_settlement}}"),
    ("(b)", "EROSION OF SLOPE:",   "{{approach_erosion}}"),
    ("(c)", "ANY OTHER SPECIFY:",  "{{approach_other}}"),
]:
    add_row(table, sr, desc, ph)

add_row(table, "7", "EXPANSION JOINT:", "{{expansion_joint}}")
add_row(table, "8", "WEARING COAT:",    "{{wearing_coat}}")

# B9 — Miscellaneous
b_section_header(table, "9", "MISCELLANEOUS")
for sr, desc, ph in [
    ("(a)", "FLOOD GAUGE MARK:",               "{{flood_gauge}}"),
    ("(b)", "MASONRY STEPS:",                  "{{masonry_steps}}"),
    ("(c)", "VEGETATION GROWTH & LITTERING:",  "{{vegetation}}"),
]:
    add_row(table, sr, desc, ph)

# ═══ SECTION C ═══════════════════════════════════════════════════════════════
section_header(table, "(C)  EVALUATION AND RECOMMENDATIONS")
add_row(table, "(i)", "CONDITION STATE AS PER IRC SP: 40-2019\n(Excellent / Good / Fair / Poor / Critical)",
        "{{condition_state}}")

# C(ii) — Recommendations header
span_row(table,
         "(ii)  RECOMMENDATIONS\n"
         "Nature of the Problem / Deficiency and Remedial Measures based on Visual Inspection:",
         bg=HEADER_BG, bold=True, size=9, color=WHITE)

# ── 1. General (Non-Structural) ───────────────────────────────────────────────
span_row(table, "1.  GENERAL (NON STRUCTURAL ELEMENTS)", bg=SECTION_BG, bold=True, size=9)

gen_items = [
    ("A)", "TRAINING & PROTECTION WORK",       "{{rec_gen_training}}"),
    ("B)", "WEARING COAT",                      "{{rec_gen_wearing}}"),
    ("C)", "VEGETATION GROWTH",                 "{{rec_gen_vegetation}}"),
    ("D)", "EXPANSION JOINT",                   "{{rec_gen_expansion}}"),
    ("E)", "MASONRY STEPS",                     "{{rec_gen_masonry}}"),
    ("F)", "FLOOD GAUGE MARK",                  "{{rec_gen_flood}}"),
    ("G)", "ANY OTHER (GENERAL)",               "{{rec_gen_other}}"),
]
for sr, desc, ph in gen_items:
    add_row(table, sr, desc, ph)

# ── 2. Structural Elements ────────────────────────────────────────────────────
span_row(table, "2.  STRUCTURAL ELEMENTS", bg=SECTION_BG, bold=True, size=9)

str_items = [
    ("A)", "SUPERSTRUCTURE\n(Girders / Deck Slab / Diaphragm)",       "{{rec_str_superstructure}}"),
    ("B)", "SUBSTRUCTURE\n(Pier / Abutment / Return Wall / Pier Cap)", "{{rec_str_substructure}}"),
    ("",   "BEARINGS",                                                  "{{rec_str_bearings}}"),
    ("C)", "FOUNDATION",                                                "{{rec_str_foundation}}"),
    ("D)", "ANY OTHER (STRUCTURAL)",                                    "{{rec_str_other}}"),
]
for sr, desc, ph in str_items:
    add_row(table, sr, desc, ph)

# ── IRC SP: 40-2019 Rating ────────────────────────────────────────────────────
span_row(table, "Recommendations based on IRC SP: 40-2019", bg=SECTION_BG, bold=True, size=9)
add_row(table, "", "Condition Rating\n(Excellent / Good / Fair / Poor / Critical)", "{{rec_irc_condition}}")
add_row(table, "", "Recommended Action",                                             "{{rec_irc_action}}")

# ── Fix column widths (python-docx doesn't always persist widths set on columns) ─
for row in table.rows:
    for i, (cell, w) in enumerate(zip(row.cells, col_widths)):
        cell.width = w

# ── Force fixed table layout so Word respects column widths ──────────────────
tblPr = table._tbl.find(qn('w:tblPr'))
if tblPr is None:
    tblPr = OxmlElement('w:tblPr')
    table._tbl.insert(0, tblPr)
tblLayout = OxmlElement('w:tblLayout')
tblLayout.set(qn('w:type'), 'fixed')
tblPr.append(tblLayout)

# ── Signature block ───────────────────────────────────────────────────────────
doc.add_paragraph()
sig = doc.add_paragraph()
sig.add_run("Sign:").bold = True
sig.add_run("  ").font.size = Pt(10)
sig.add_run("Date: ").bold = True
sig.add_run("{{sign_date}}")

for label, placeholder in [
    ("1) Signature of Consultant:",  "_______________________"),
    ("2) Name of Representative:",   "{{representative_name}}"),
    ("3) Name of Consultant's Firm:","CASAD CONSULTANTS PVT. LTD."),
    ("4) Signature of Field Officer:","_______________________"),
]:
    p = doc.add_paragraph()
    r = p.add_run(f"{label}  ")
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(placeholder).font.size = Pt(10)

# ── Disclaimer page ───────────────────────────────────────────────────────────
doc.add_page_break()
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
rh = h.add_run("Basis and Limitations of Visual Inspection")
rh.bold = True
rh.font.size = Pt(12)

DISCLAIMER = (
    "This report has been prepared based solely on visual inspection of the bridge "
    "structures carried out by our team with due care, diligence, and professional "
    "judgment. All distresses, defects, and observable damages that were visible and "
    "accessible at the time of inspection have been recorded to the best of our "
    "knowledge and technical competence.\n\n"
    "It is clearly understood that this assessment is limited to surface-level and "
    "visible conditions only. Any hidden, internal, or latent defects within structural "
    "members, foundations, substructure, or materials that are not apparent during "
    "visual inspection are beyond the scope of this report and could not be identified.\n\n"
    "This report does not constitute a guarantee or warranty of the present or future "
    "structural performance, safety, or serviceability of the bridge. The behavior of "
    "the structure over time depends on several factors including loading, ageing, "
    "environmental exposure, maintenance practices, and unforeseen external actions, "
    "which are beyond our control.\n\n"
    "Accordingly, our firm and inspection team shall not be held liable for any future "
    "deterioration, distress, failure, or damage that may occur after the date of "
    "inspection, nor for any consequences arising from such future behavior of the "
    "the structure."
)
dp = doc.add_paragraph(DISCLAIMER)
dp.paragraph_format.space_before = Pt(10)
dp.runs[0].font.size = Pt(10)

# ── Appendix A — General Site Pictures ───────────────────────────────────────
doc.add_page_break()
h_a = doc.add_paragraph()
h_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_a = h_a.add_run("Appendix A: General Site Pictures")
r_a.bold = True
r_a.font.size = Pt(12)
r_a.font.color.rgb = RGBColor.from_string(DARK_TEXT)

marker_a = doc.add_paragraph("[[PHOTO_APPENDIX_A]]")
marker_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
marker_a.runs[0].font.color.rgb = RGBColor.from_string("AAAAAA")
marker_a.runs[0].font.size = Pt(9)
marker_a.runs[0].italic = True

# ── Appendix B — Damage / Distressing Photographs ────────────────────────────
doc.add_page_break()
h_b = doc.add_paragraph()
h_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_b = h_b.add_run("Appendix B: Damage / Distressing Photographs")
r_b.bold = True
r_b.font.size = Pt(12)
r_b.font.color.rgb = RGBColor.from_string(DARK_TEXT)

marker_b = doc.add_paragraph("[[PHOTO_APPENDIX_B]]")
marker_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
marker_b.runs[0].font.color.rgb = RGBColor.from_string("AAAAAA")
marker_b.runs[0].font.size = Pt(9)
marker_b.runs[0].italic = True

out = "casad_template.docx"
doc.save(out)
print(f"Template saved: {out}")
