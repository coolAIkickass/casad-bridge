# checker.py — Post-generation report quality checker and corrector
#
# check_report()   → detects issues without modifying any file
# correct_report() → fixes auto_fixable issues in-place, then re-checks once
# log_issues()     → appends structured entries to checker_errors.log
#
# All cell addresses and field names are sourced verbatim from
# report_gen_excel.py, report_gen_excel_amc.py, and report_gen.py.

import os
import io
from datetime import datetime, timezone, date as _date_type
from dataclasses import dataclass, field
from typing import List, Optional, Set, Any

LOG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'checker_errors.log')

SPECIAL_VALUES = frozenset([
    "Not Applicable",
    "Data Not Available",
    "Absent",
    "Not Visible",
    "As per approved GAD",
])


# ─────────────────────────────────────────────────────────────────────────────
#  Data classes
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class CheckIssue:
    severity: str        # 'warning' | 'error'
    rule: int            # 1–4
    cell_or_field: str
    found: str
    expected: str
    auto_fixable: bool
    was_corrected: bool = False


@dataclass
class CheckResult:
    issues: List[CheckIssue] = field(default_factory=list)

    @property
    def has_issues(self) -> bool:
        return bool(self.issues)


# ─────────────────────────────────────────────────────────────────────────────
#  Low-level helpers
# ─────────────────────────────────────────────────────────────────────────────

def _vj(d: dict, key: str):
    """Replicate the generator's inner _v() closure: return None for missing/dash values."""
    v = d.get(key)
    return v if v and str(v).strip() not in ('-', '') else None


def _fmt_date_safe(val) -> Optional[str]:
    """Format date field as dd/mm/yyyy string, or None if empty."""
    if val is None or str(val).strip() in ('', '-'):
        return None
    if isinstance(val, (_date_type, datetime)):
        return val.strftime('%d/%m/%Y')
    for fmt in ('%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y'):
        try:
            return datetime.strptime(str(val), fmt).strftime('%d/%m/%Y')
        except (ValueError, TypeError):
            pass
    s = str(val).strip()
    return s if s else None


def _str_val(v) -> Optional[str]:
    """Convert any cell value (including CellRichText, datetime) to plain str or None."""
    if v is None:
        return None
    try:
        from openpyxl.cell.rich_text import CellRichText
        if isinstance(v, CellRichText):
            s = str(v).strip()
            return s if s else None
    except ImportError:
        pass
    if isinstance(v, (_date_type, datetime)):
        return v.strftime('%d/%m/%Y')
    s = str(v).strip()
    return s if s not in ('-', '', 'None') else None


def _read_cell_val(ws, addr: str):
    """Read cell value following merge ranges to find the anchor cell."""
    from openpyxl.utils import coordinate_to_tuple, range_boundaries
    from openpyxl.cell.cell import MergedCell
    row, col = coordinate_to_tuple(addr)
    cell = ws.cell(row=row, column=col)
    if isinstance(cell, MergedCell):
        for rng in ws.merged_cells.ranges:
            mc, mr, xc, xr = range_boundaries(str(rng))
            if mr <= row <= xr and mc <= col <= xc:
                return ws.cell(row=mr, column=mc).value
        return None
    return cell.value


def _write_cell(ws, addr: str, value):
    """Write to cell, unmerging if it is a MergedCell."""
    from openpyxl.utils import coordinate_to_tuple, range_boundaries
    from openpyxl.cell.cell import MergedCell
    row, col = coordinate_to_tuple(addr)
    cell = ws.cell(row=row, column=col)
    if isinstance(cell, MergedCell):
        for rng in list(ws.merged_cells.ranges):
            mc, mr, xc, xr = range_boundaries(str(rng))
            if mr <= row <= xr and mc <= col <= xc:
                ws.unmerge_cells(str(rng))
                break
    ws.cell(row=row, column=col).value = value


def _find_ws(wb, name_or_keyword: str, exact: bool = True):
    """Find a worksheet by exact name (exact=True) or keyword (exact=False)."""
    if exact:
        return wb[name_or_keyword] if name_or_keyword in wb.sheetnames else None
    kw = name_or_keyword.lower()
    for sn in wb.sheetnames:
        if kw in sn.lower():
            return wb[sn]
    return None


# ─────────────────────────────────────────────────────────────────────────────
#  Traceable-value set builder
# ─────────────────────────────────────────────────────────────────────────────

def _build_value_set(d: dict) -> Set[str]:
    """Return a set of every string value that is legitimately traceable to report_json."""
    values: Set[str] = set()

    def _collect(v):
        if v is None:
            return
        if isinstance(v, str):
            s = v.strip()
            if s and s not in ('-', ''):
                values.add(s)
        elif isinstance(v, (int, float)):
            values.add(str(v))
        elif isinstance(v, list):
            for item in v:
                _collect(item)
        elif isinstance(v, dict):
            for val in v.values():
                _collect(val)

    for v in d.values():
        _collect(v)

    # Computed combinations that the generator writes as single cell strings
    lat = _vj(d, 'latitude')
    lon = _vj(d, 'longitude')
    if lat and lon:
        values.add(f"{lat}° , {lon}°")

    div = _vj(d, 'division')
    cir = _vj(d, 'circle')
    if div and cir and div != cir:
        values.add(f"{div} / {cir}")

    # C15 (Appendix-A): _build_spans_cell wraps no_of_spans in "Number of Span: …"
    nos = _vj(d, 'no_of_spans')
    if nos:
        values.add('Number of Span: ' + str(nos))

    # Date normalisation: add dd/mm/yyyy form for any raw date field
    for key in ('date_of_survey', 'date_of_construction_start', 'date_of_completion'):
        raw = d.get(key)
        if raw:
            formatted = _fmt_date_safe(raw)
            if formatted:
                values.add(formatted)

    return values


# ─────────────────────────────────────────────────────────────────────────────
#  Expected-value functions (replicate generator logic exactly)
# ─────────────────────────────────────────────────────────────────────────────

def _appendix_a_expected(addr: str, d: dict):
    """
    Return the value the generator would write to Appendix-A column-C cell `addr`.
    Replicates _fill_appendix_a() from both report_gen_excel.py and
    report_gen_excel_amc.py (the two are identical for this sheet).
    """
    def _v(k):
        return _vj(d, k)

    lat = _v('latitude')
    lon = _v('longitude')
    location = f"{lat}° , {lon}°" if lat and lon else None

    # _build_spans_cell produces CellRichText; fall back to plain string if import fails
    try:
        from report_gen_excel import _build_spans_cell
        spans_val = _build_spans_cell(d)
    except Exception:
        nos = _v('no_of_spans')
        spans_val = ('Number of Span: ' + nos) if nos else None

    m = {
        # Section 1 — General identity
        'C4':  _v('bridge_title') or _v('river_name'),
        'C5':  _v('bridge_number'),
        'C6':  _v('river_name'),
        'C7':  _v('road_name'),
        'C8':  _v('road_number'),
        'C9':  location,
        'C10': _v('bm_gts_level'),
        'C11': _v('division') or _v('circle'),
        'C12': _v('circle'),
        # Section 2 — Details of Spans
        'C15': spans_val,
        'C16': _v('total_length'),
        'C17': _v('angle_of_crossing'),
        'C18': _v('bridge_level_type') or _v('type_of_bridge'),
        # Section 3 — Hydraulic Parameters
        'C20': _v('hydraulic_parameters'),
        'C21': _v('hydraulic_catchment'),
        'C22': _v('hydraulic_discharge'),
        'C23': _v('hydraulic_hfl'),
        'C24': _v('hydraulic_ofl'),
        'C25': _v('hydraulic_clearance'),
        'C26': _v('hydraulic_lwl'),
        'C27': _v('hydraulic_depth'),
        'C28': _v('hydraulic_velocity'),
        'C29': _v('hydraulic_channel_width'),
        'C30': _v('hydraulic_spread'),
        'C31': _v('hydraulic_bed_level'),
        # Section 4 — Sub Soil Particulars
        'C33': _v('subsoil_particulars'),
        'C34': _v('subsoil_type'),
        'C35': _v('subsoil_friction'),
        'C36': _v('subsoil_cohesion'),
        'C37': _v('subsoil_silt_factor'),
        'C38': _v('subsoil_bearing_capacity'),
        'C39': _v('subsoil_foundation_level'),
        # Section 5 — Design and Structural Data
        'C42': _v('loading_standard'),
        'C43': _v('superstructure_type') or _v('bridge_type'),
        'C44': _v('span_arrangement') or _v('span_length'),
        'C45': _v('carriage_width'),
        'C46': _v('deck_level'),
        'C48': _v('design_scour_level'),
        'C49': _v('design_foundation_level'),
        'C50': _v('foundation_type'),
        'C51': _v('substructure_type'),
        'C52': _v('substructure_material'),
        'C53': _v('pier_length'),
        'C54': _v('pier_width_detail'),
        'C55': _v('pier_cap_width'),
        'C56': _v('abutment_width'),
        'C57': _v('abutment_cap_width'),
        'C58': _v('returns_length'),
        'C59': _v('superstructure_type'),
        'C60': _v('prestressing_details'),
        'C61': _v('articulation_details'),
        'C62': _v('total_load_foundation'),
        'C63': _v('total_horizontal_force'),
        'C64': _v('bearing_type_detail'),
        'C65': _v('wearing_coat'),
        'C66': _v('railing_type'),
        'C67': _v('expansion_joint'),
        'C68': _v('protection_works'),
        'C69': _v('model_studies'),
        'C70': _v('special_design_features'),
        'C71': _v('settlement_report'),
        # Material Consumed
        'C73': _v('material_consumed'),
        'C74': _v('material_cement'),
        'C75': _v('material_reinforcement'),
        'C76': _v('material_structural_steel'),
        'C77': _v('material_hts_steel'),
        # Other Data
        'C80': _fmt_date_safe(d.get('date_of_construction_start')),
        'C81': _fmt_date_safe(d.get('date_of_completion')),
        'C82': _v('surface_utilities'),
        'C83': _v('design_drawings'),
        'C84': _v('ls_sketch'),
        'C85': _v('special_features'),
        'C86': _v('total_cost'),
        'C87': _v('cost_per_sqm_carriageway'),
        'C88': _v('cost_per_sqm_elevation'),
        'C89': _v('cost_per_m_length'),
        'C90': _v('design_agency'),
        'C91': _v('construction_agency'),
        # Performance & survey date
        'C92': _v('performance'),
        'C93': _fmt_date_safe(d.get('date_of_survey')),
    }
    return m.get(addr)


def _appendix_b_expected(addr: str, d: dict, fmt: str):
    """
    Return the value the generator would write to Appendix-B column-C cell `addr`.
    Covers both 'excel_rb' (report_gen_excel._fill_appendix_b) and
    'excel_amc' (report_gen_excel_amc._fill_appendix_b).
    """
    def _v(k):
        return _vj(d, k)

    lat = _v('latitude')
    lon = _v('longitude')
    location = f"{lat}° , {lon}°" if lat and lon else None

    div = _v('division')
    cir = _v('circle')
    if div and cir and div != cir:
        div_cir = f"{div} / {cir}"
    else:
        div_cir = div or cir or None

    # Common cells (identical in both R&B and AMC _fill_appendix_b)
    m = {
        # Section 1 — General
        'C4':  _v('bridge_title') or _v('river_name'),
        'C5':  _v('bridge_number'),
        'C6':  _v('river_name'),
        'C7':  _v('road_name'),
        'C8':  _v('road_number'),
        'C9':  location,
        'C10': div_cir,
        'C11': _v('type_of_bridge') or _v('bridge_type'),
        'C12': _fmt_date_safe(d.get('date_of_survey')),
        # Section 4 — Approaches
        'C14': _v('approach_settlement'),
        'C15': _v('approach_side_slopes'),
        'C16': _v('approach_erosion'),
        'C17': _v('approach_slab'),
        'C18': _v('approach_geometrics'),
        'C19': _v('approach_other'),
        # Section 5 — Protective Works
        'C21': _v('prot_type'),
        'C22': _v('prot_damage_layout'),
        'C23': _v('prot_slope_pitching'),
        'C24': _v('prot_floor_protection'),
        'C25': _v('prot_scour_extent'),
        'C26': _v('prot_reserve_stone'),
        'C27': _v('prot_other'),
        # Section 6 — Waterway
        'C28': _v('waterway_obs'),
        'C29': _v('waterway_obstruction'),
        'C30': _v('waterway_scour'),
        'C31': _v('waterway_flow'),
        'C32': _v('waterway_flood_level'),
        'C33': _v('waterway_afflux'),
        'C34': _v('waterway_adequacy'),
        'C35': _v('waterway_other'),
        # Section 7 — Foundations
        'C36': _v('foundations_obs'),
        'C37': _v('foundations_settlement'),
        'C38': _v('foundations_cracking'),
        'C39': _v('foundations_floating'),
        'C40': _v('foundations_subway'),
        'C41': _v('foundations_other'),
        # Section 8 — Substructure
        'C42': _v('sub_section_obs'),
        'C43': _v('sub_drainage_backfill'),
        'C44': _v('sub_cracking_obs'),
        'C45': _v('sub_subway_obs'),
        'C46': _v('sub_other_obs'),
        # Section 9 — Bearings
        'C48': _v('bear_metallic_type'),
        'C49': _v('bear_metallic_condition'),
        'C50': _v('bear_metallic_functioning'),
        'C51': _v('bear_metallic_greasing'),
        'C52': _v('bear_pedestal_cracks'),
        'C53': _v('bear_metallic_anchor'),
        'C54': _v('bear_metallic_other'),
        'C55': _v('bear_elastomeric_type'),
        'C56': _v('bear_pad_condition'),
        'C57': _v('bear_cleanliness'),
        'C58': _v('bear_elastomeric_other'),
        # Section 10 — Superstructure
        'C59': _v('super_section_obs') or _v('superstructure_type'),
        'C60': None,  # intentionally blank (row 10.1 sub-section label)
        'C61': _v('super_spalling_obs'),
        'C62': _v('super_cracking_obs'),
        'C63': _v('super_corrosion_obs'),
        'C64': _v('super_vehicle_damage'),
        'C65': _v('super_articulation'),
        'C66': _v('super_vibration'),
        'C67': _v('super_deflection'),
        'C68': _v('super_anchorage_cracks'),
        'C69': _v('super_hinge_deflection'),
        'C70': _v('steel_obs'),
        'C79': _v('masonry_obs'),
        'C86': _v('timber_obs'),
        # Section 11 — Expansion Joints
        'C92': _v('exp_jt_functioning'),
        'C93': _v('exp_jt_sealing'),
        'C94': _v('exp_jt_fixing'),
        'C95': _v('exp_jt_sliding_plate'),
        'C96': _v('exp_jt_locking'),
        'C97': _v('exp_jt_debris'),
        'C98': _v('exp_jt_rattling'),
        'C99': _v('exp_jt_other'),
        # Section 12 — Wearing Coat
        'C100': _v('wear_coat_type'),
        'C101': _v('wear_coat_surface'),
        'C102': _v('wear_coat_evidence'),
        # Section 13 — Drainage Spouts
        'C103': _v('drain_type'),
        'C104': _v('drain_clogging'),
        'C105': _v('drain_projection'),
        'C106': _v('drain_adequacy'),
        'C107': _v('drain_subway'),
        'C108': _v('drain_other'),
        # Section 14 — Handrail
        'C110': _v('handrail_condition'),
        'C111': _v('handrail_collision'),
        'C112': _v('handrail_alignment'),
        # Section 15 — Footpath
        'C114': _v('footpath_condition'),
        'C115': _v('footpath_missing_slab'),
        'C116': _v('footpath_other'),
        # Section 16 — Utilities
        'C117': _v('utilities_obs'),
        'C118': _v('util_water_leakage'),
        'C119': _v('util_cable_damage'),
        'C120': _v('util_lighting'),
        'C121': _v('util_other_damage'),
        # Section 17–19
        'C123': _v('bridge_num_condition'),
        'C125': _v('aesthetics_intrusion'),
        'C126': _v('maintenance_history'),
    }

    # AMC-only cells: steel sub-rows (C71–C78), masonry sub-rows (C80–C85),
    # timber sub-rows (C87–C90).  R&B skips these.
    if fmt == 'excel_amc':
        m.update({
            'C71': _v('steel_paint'),
            'C72': _v('steel_corrosion'),
            'C73': _v('steel_vibration'),
            'C74': _v('steel_alignment'),
            'C75': _v('steel_connections'),
            'C76': _v('steel_camber_deflection'),
            'C77': _v('steel_buckling'),
            'C78': _v('steel_cleanliness'),
            'C80': _v('masonry_joints'),
            'C81': _v('masonry_profile'),
            'C82': _v('masonry_cracks'),
            'C83': _v('masonry_drainage'),
            'C84': _v('masonry_vegetation'),
            'C85': _v('masonry_other'),
            'C87': _v('timber_paint'),
            'C88': _v('timber_decay'),
            'C89': _v('timber_joints'),
            'C90': _v('timber_sag'),
            # AMC overall condition is at row 127
            'C127': _v('overall_condition_visual'),
        })
    else:
        # R&B overall condition is at row 128
        m['C128'] = _v('overall_condition_visual')

    return m.get(addr)


# ─────────────────────────────────────────────────────────────────────────────
#  Rule 5 — Raw input validation (raw_bridge_details → report_json)
# ─────────────────────────────────────────────────────────────────────────────

# (label pattern lowercase, field_name, Appendix-A cell addr)
LABEL_FIELD_MAP = [
    # Section 2 — Spans / geometry
    ('span arrangement',               'span_arrangement',          'C44'),
    ('angle of crossing',              'angle_of_crossing',         'C17'),
    # Section 5 — Design data (commonly missed because not in STYLE 3 mapping)
    ('designed maximum scour level',   'design_scour_level',        'C48'),
    ('design scour level',             'design_scour_level',        'C48'),
    ('maximum scour level',            'design_scour_level',        'C48'),
    ('designed foundation level',      'design_foundation_level',   'C49'),
    ('design foundation level',        'design_foundation_level',   'C49'),
    ('foundation level from pile cap', 'design_foundation_level',   'C49'),
    # Substructure dimensions
    ('straight length of pier',        'pier_length',               'C53'),
    ('pier length',                    'pier_length',               'C53'),
    ('width of pier gap',              'pier_cap_width',            'C55'),
    ('pier gap',                       'pier_cap_width',            'C55'),
    ('width of pier cap',              'pier_cap_width',            'C55'),
    ('pier cap',                       'pier_cap_width',            'C55'),
    ('width of abutment cap',          'abutment_cap_width',        'C57'),
    ('abutment cap',                   'abutment_cap_width',        'C57'),
    ('width of abutment',              'abutment_width',            'C56'),
    # Superstructure construction details (sub-section j)
    ('details of prestressing',        'prestressing_details',      'C60'),
    ('details of articulation',        'articulation_details',      'C61'),
    # Agencies / dates
    ('design agency',                  'design_agency',             'C90'),
    ('construction agency',            'construction_agency',       'C91'),
    ('date of completion',             'date_of_completion',        'C81'),
    ('date of construction',           'date_of_construction_start','C80'),
]

_NORMALIZE = {
    'not applicable': 'Not Applicable',
    'n/a': 'Not Applicable',
    'data not available': 'Data Not Available',
    'not available': 'Data Not Available',
    'same as above': 'Same as above',
    'absent': 'Absent',
    'not visible': 'Not Visible',
    'as per approved gad': 'As per approved GAD',
}
_STOP_WORDS = ['next ', ' and then ', '\n']


def _extract_value_after_label(raw_text: str, pattern: str) -> Optional[str]:
    """Extract the value following a label pattern in raw text.

    Conservative: only returns when the result is unambiguous (special phrase
    or simple numeric measurement). Returns None when uncertain.
    """
    import re as _re
    lower = raw_text.lower()
    idx = lower.find(pattern)
    if idx == -1:
        return None
    remainder = raw_text[idx + len(pattern):].strip()
    # Trim at next recognised stop boundary
    end = len(remainder)
    for stop in _STOP_WORDS:
        p = remainder.lower().find(stop)
        if 0 < p < end:
            end = p
    chunk = remainder[:end].strip().rstrip('.,')
    if not chunk:
        return None
    # Match against normalised special values
    chunk_lower = chunk.lower()
    for key, val in _NORMALIZE.items():
        if chunk_lower.startswith(key):
            return val
    # Match a simple measurement: number + optional unit
    m = _re.match(r'^(\d+(?:\.\d+)?)\s*(m|meter|metre|mtr)?\b', chunk, _re.IGNORECASE)
    if m:
        num = m.group(1)
        unit = (m.group(2) or '').lower()
        return f"{num} m" if unit in ('m', 'meter', 'metre', 'mtr', '') else num
    # Free-text: return as-is only if short enough to be unambiguous
    return chunk if len(chunk) <= 80 else None


def _check_raw_input(raw_text: str, d: dict, result: CheckResult) -> None:
    """Rule 5: labels found in raw bridge_details text must have non-null JSON values."""
    import re as _re
    lower = raw_text.lower()
    seen_fields: Set[str] = set()
    for pattern, field, cell in LABEL_FIELD_MAP:
        if field in seen_fields:
            continue
        if pattern not in lower:
            continue
        seen_fields.add(field)
        if _vj(d, field) is not None:
            continue  # already populated — no issue
        extracted = _extract_value_after_label(raw_text, pattern)
        fixable = extracted is not None and (
            extracted in _NORMALIZE.values()
            or bool(_re.match(r'^\d', extracted))
        )
        result.issues.append(CheckIssue(
            severity='warning',
            rule=5,
            cell_or_field=f'Appendix-A:{cell}',
            found='(blank in report_json)',
            expected=extracted or f'(value after label "{pattern}")',
            auto_fixable=fixable,
        ))


# ─────────────────────────────────────────────────────────────────────────────
#  Rule 1 — Lost inputs  (Excel)
# ─────────────────────────────────────────────────────────────────────────────

def _check_rule1_excel(wb, d: dict, fmt: str, result: CheckResult) -> None:
    """Check 1: every non-empty report_json value must appear in its expected cell."""
    ws_a = _find_ws(wb, 'Appendix-A')
    ws_b = _find_ws(wb, 'Appendix-B')

    # Appendix-A: column C rows 4–93 (last mapped row is C93 = date_of_survey)
    if ws_a:
        for row in range(4, 94):
            addr = f'C{row}'
            expected = _appendix_a_expected(addr, d)
            if expected is None:
                continue
            exp_str = _str_val(expected)
            if exp_str is None:
                continue
            actual_str = _str_val(_read_cell_val(ws_a, addr))
            if actual_str is None:
                result.issues.append(CheckIssue(
                    severity='warning', rule=1,
                    cell_or_field=f'Appendix-A:{addr}',
                    found='(empty)',
                    expected=exp_str[:120],
                    auto_fixable=True,
                ))

    # Appendix-B: column C rows 4–128 (R&B) or 4–127 (AMC)
    if ws_b:
        max_row = 127 if fmt == 'excel_amc' else 128
        for row in range(4, max_row + 1):
            addr = f'C{row}'
            expected = _appendix_b_expected(addr, d, fmt)
            if expected is None:
                continue
            exp_str = _str_val(expected)
            if exp_str is None:
                continue
            actual_str = _str_val(_read_cell_val(ws_b, addr))
            if actual_str is None:
                result.issues.append(CheckIssue(
                    severity='warning', rule=1,
                    cell_or_field=f'Appendix-B:{addr}',
                    found='(empty)',
                    expected=exp_str[:120],
                    auto_fixable=True,
                ))


# ─────────────────────────────────────────────────────────────────────────────
#  Rule 2 — Hallucinations in protected sections  (Excel)
# ─────────────────────────────────────────────────────────────────────────────

def _check_rule2_excel(wb, d: dict, fmt: str, result: CheckResult) -> None:
    """Check 2: every non-blank cell in Appendix-A/B col-C must trace to report_json."""
    traceable = _build_value_set(d)
    ws_a = _find_ws(wb, 'Appendix-A')
    ws_b = _find_ws(wb, 'Appendix-B')

    def _scan(ws, sheet_tag, row_start, row_end):
        for row in range(row_start, row_end + 1):
            addr = f'C{row}'
            raw = _read_cell_val(ws, addr)
            cell_str = _str_val(raw)
            if cell_str is None:
                continue
            # A cell is traceable if its value (or a substring) is in our set,
            # or if the traceable set contains a substring of the cell value
            # (covers cases like "Elastomeric Bearing (Side A)" ⊃ "Elastomeric Bearing").
            is_traceable = (
                cell_str in traceable
                or any(cell_str in tv or tv in cell_str for tv in traceable if len(tv) >= 3)
            )
            if not is_traceable:
                result.issues.append(CheckIssue(
                    severity='warning', rule=2,
                    cell_or_field=f'{sheet_tag}:{addr}',
                    found=cell_str[:120],
                    expected='value traceable to report_json',
                    auto_fixable=True,
                ))

    if ws_a:
        _scan(ws_a, 'Appendix-A', 4, 104)
    if ws_b:
        _scan(ws_b, 'Appendix-B', 4, 129)


# ─────────────────────────────────────────────────────────────────────────────
#  Rule 3 — Stretched photos  (Excel)
# ─────────────────────────────────────────────────────────────────────────────

def _check_rule3_excel(wb, d: dict, fmt: str, result: CheckResult) -> None:
    """Check 3: embedded images must not have their aspect ratio distorted > 25%."""
    photos = d.get('photos', [])
    cats   = d.get('photo_categories', [])
    if not photos:
        return

    try:
        from PIL import Image as PILImage
    except ImportError:
        return  # PIL unavailable — skip photo checks

    def _check_sheet_photos(ws, source_paths):
        if ws is None or not source_paths:
            return
        images = getattr(ws, '_images', [])
        for idx, img in enumerate(images):
            if idx >= len(source_paths):
                break
            src_path = source_paths[idx]
            if not src_path or not os.path.exists(src_path):
                continue
            try:
                with PILImage.open(src_path) as src_img:
                    src_w, src_h = src_img.size
                if src_w == 0 or src_h == 0:
                    continue
                src_aspect = src_w / src_h
                disp_w = getattr(img, 'width',  None)
                disp_h = getattr(img, 'height', None)
                if disp_w and disp_h and disp_h != 0:
                    disp_aspect = disp_w / disp_h
                    deviation = abs(disp_aspect - src_aspect) / src_aspect
                    if deviation > 0.25:
                        sheet_name = getattr(ws, 'title', 'unknown')
                        result.issues.append(CheckIssue(
                            severity='warning', rule=3,
                            cell_or_field=f'{sheet_name}:image[{idx}]',
                            found=f'display {disp_w}×{disp_h} (aspect {disp_aspect:.3f})',
                            expected=f'source aspect {src_aspect:.3f} (±25%)',
                            auto_fixable=True,
                        ))
            except Exception:
                pass

    if fmt == 'excel_rb':
        sub_paths   = [p for p, c in zip(photos, cats) if c == 'general']
        super_paths = [p for p, c in zip(photos, cats) if c == 'damaged']
        ws_sub   = _find_ws(wb, 'Appendix__C')
        ws_super = _find_ws(wb, 'Appendix__C (2)')
        _check_sheet_photos(ws_sub,   sub_paths)
        _check_sheet_photos(ws_super, super_paths)
    elif fmt == 'excel_amc':
        ws_photo = _find_ws(wb, 'appendix - photos', exact=False)
        _check_sheet_photos(ws_photo, photos)


# ─────────────────────────────────────────────────────────────────────────────
#  Rule 4 — Special values preserved  (Excel)
# ─────────────────────────────────────────────────────────────────────────────

def _check_rule4_excel(wb, d: dict, fmt: str, result: CheckResult) -> None:
    """Check 4: special user-entered phrases must survive generation unchanged."""
    ws_a = _find_ws(wb, 'Appendix-A')
    ws_b = _find_ws(wb, 'Appendix-B')

    def _check(ws, sheet_tag, expected_fn, max_row):
        if ws is None:
            return
        for row in range(4, max_row + 1):
            addr = f'C{row}'
            exp = expected_fn(addr)
            if exp is None:
                continue
            exp_str = _str_val(exp)
            if exp_str not in SPECIAL_VALUES:
                continue
            actual_str = _str_val(_read_cell_val(ws, addr))
            if actual_str != exp_str:
                result.issues.append(CheckIssue(
                    severity='error', rule=4,
                    cell_or_field=f'{sheet_tag}:{addr}',
                    found=actual_str or '(empty)',
                    expected=exp_str,
                    auto_fixable=True,
                ))

    if ws_a:
        _check(ws_a, 'Appendix-A', lambda addr: _appendix_a_expected(addr, d), 93)
    if ws_b:
        max_row = 127 if fmt == 'excel_amc' else 128
        _check(ws_b, 'Appendix-B', lambda addr: _appendix_b_expected(addr, d, fmt), max_row)


# ─────────────────────────────────────────────────────────────────────────────
#  Word checks (Rules 1 and 4)
# ─────────────────────────────────────────────────────────────────────────────

_WORD_SECTION_A_FIELDS = (
    'river_name', 'road_name', 'chainage', 'latitude', 'longitude',
    'circle', 'division', 'sub_division', 'no_of_spans', 'span_length',
    'span_arrangement', 'total_length', 'bridge_type', 'bridge_level_type',
    'type_of_bridge', 'superstructure_type', 'substructure_type',
    'foundation_type', 'bearing_type_detail', 'approach_length',
    'railing_type', 'river_training', 'repair_work', 'carriage_width',
    'year_of_construction', 'river_perennial', 'angle_of_crossing',
    'hydraulic_parameters', 'subsoil_particulars', 'deck_level',
    'pier_length', 'date_of_completion', 'surface_utilities', 'performance',
    'prestressing_details', 'date_of_survey',
)


def _check_rule1_word(doc, d: dict, result: CheckResult) -> None:
    """Check 1 (Word): Section A field values must appear in the document body."""
    # Extract all text from paragraphs and table cells
    all_text_parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_text_parts.append(para.text)
    full_text = '\n'.join(all_text_parts)

    for key in _WORD_SECTION_A_FIELDS:
        val = _vj(d, key)
        if val is None:
            continue
        val_str = str(val).strip()
        if not val_str or val_str in ('-', ''):
            continue
        if val_str not in full_text:
            result.issues.append(CheckIssue(
                severity='warning', rule=1,
                cell_or_field=f'word:section_a:{key}',
                found='(absent from document)',
                expected=val_str[:120],
                auto_fixable=False,  # Word placeholder replacement is already done
            ))


def _check_rule4_word(doc, d: dict, result: CheckResult) -> None:
    """Check 4 (Word): special phrases from report_json must appear in the document."""
    all_text_parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_text_parts.append(para.text)
    full_text = '\n'.join(all_text_parts)

    for key, val in d.items():
        if not isinstance(val, str):
            continue
        val_str = val.strip()
        if val_str not in SPECIAL_VALUES:
            continue
        if val_str not in full_text:
            result.issues.append(CheckIssue(
                severity='error', rule=4,
                cell_or_field=f'word:{key}',
                found='(absent from document)',
                expected=val_str,
                auto_fixable=False,
            ))


# ─────────────────────────────────────────────────────────────────────────────
#  check_report() entry point
# ─────────────────────────────────────────────────────────────────────────────

def check_report(report_path: str, report_json: dict, fmt: str,
                 raw_bridge_details: Optional[str] = None) -> CheckResult:
    """
    Run all quality checks on the generated report file.

    Parameters
    ----------
    report_path         : str            path to the generated .xlsx or .docx
    report_json         : dict           structured JSON produced by ai_parse.py
    fmt                 : str            'excel_rb' | 'excel_amc' | 'word'
    raw_bridge_details  : str | None     joined text of all bridge_details messages;
                                         when provided, Rule 5 (raw input validation)
                                         is also run

    Returns
    -------
    CheckResult  containing a list of CheckIssue objects
    """
    result = CheckResult()
    try:
        if fmt in ('excel_rb', 'excel_amc'):
            import openpyxl
            wb = openpyxl.load_workbook(report_path)
            try:
                _check_rule1_excel(wb, report_json, fmt, result)
                _check_rule2_excel(wb, report_json, fmt, result)
                _check_rule3_excel(wb, report_json, fmt, result)
                _check_rule4_excel(wb, report_json, fmt, result)
            finally:
                wb.close()
        else:
            from docx import Document
            doc = Document(report_path)
            _check_rule1_word(doc, report_json, result)
            _check_rule4_word(doc, report_json, result)

        # Rule 5 — raw input validation (Excel only; independent of file open)
        if raw_bridge_details and fmt in ('excel_rb', 'excel_amc'):
            _check_raw_input(raw_bridge_details, report_json, result)

    except Exception as exc:
        print(f'[CHECKER] check_report failed: {exc}', flush=True)

    return result


# ─────────────────────────────────────────────────────────────────────────────
#  Correction helpers (Excel)
# ─────────────────────────────────────────────────────────────────────────────

def _correct_excel(report_path: str, d: dict, fmt: str, issues: List[CheckIssue]) -> None:
    """Apply auto-fixable corrections to the Excel report in-place."""
    import openpyxl
    wb = openpyxl.load_workbook(report_path)
    modified = False

    for issue in issues:
        if not issue.auto_fixable:
            continue
        try:
            parts = issue.cell_or_field.split(':')
            if len(parts) != 2:
                continue
            sheet_tag, addr = parts

            # ── Rule 1: write missing value ──────────────────────────────────
            if issue.rule == 1:
                ws = None
                if sheet_tag == 'Appendix-A':
                    ws = _find_ws(wb, 'Appendix-A')
                    expected = _appendix_a_expected(addr, d)
                elif sheet_tag == 'Appendix-B':
                    ws = _find_ws(wb, 'Appendix-B')
                    expected = _appendix_b_expected(addr, d, fmt)
                else:
                    continue
                if ws is not None and expected is not None:
                    _write_cell(ws, addr, expected)
                    issue.was_corrected = True
                    modified = True

            # ── Rule 2: clear hallucinated cell ─────────────────────────────
            elif issue.rule == 2:
                ws = None
                if sheet_tag == 'Appendix-A':
                    ws = _find_ws(wb, 'Appendix-A')
                elif sheet_tag == 'Appendix-B':
                    ws = _find_ws(wb, 'Appendix-B')
                if ws is not None:
                    _write_cell(ws, addr, None)
                    issue.was_corrected = True
                    modified = True

            # ── Rule 3: re-embed stretched photo ────────────────────────────
            elif issue.rule == 3:
                # addr is 'image[N]' — parse sheet name from sheet_tag
                # sheet_tag here is the actual sheet name (e.g. 'Appendix__C')
                ws = _find_ws(wb, sheet_tag) or _find_ws(wb, sheet_tag, exact=False)
                if ws is None:
                    continue
                # Parse image index from cell_or_field suffix 'image[N]'
                import re as _re
                m = _re.search(r'image\[(\d+)\]', addr)
                if not m:
                    continue
                img_idx = int(m.group(1))
                images = getattr(ws, '_images', [])
                if img_idx >= len(images):
                    continue

                # Find source path
                photos = d.get('photos', [])
                cats   = d.get('photo_categories', [])
                if fmt == 'excel_rb':
                    if 'Appendix__C (2)' in sheet_tag or 'appendix__c (2)' in sheet_tag.lower():
                        src_paths = [p for p, c in zip(photos, cats) if c == 'damaged']
                    else:
                        src_paths = [p for p, c in zip(photos, cats) if c == 'general']
                else:
                    src_paths = list(photos)

                if img_idx >= len(src_paths):
                    continue
                src_path = src_paths[img_idx]
                if not src_path or not os.path.exists(src_path):
                    continue

                try:
                    from PIL import Image as PILImage
                    from openpyxl.drawing.image import Image as XLImage
                    old_img = images[img_idx]
                    anchor  = str(old_img.anchor)
                    ws._images.remove(old_img)

                    with PILImage.open(src_path) as src_img:
                        src_img.load()
                        if src_img.mode in ('RGBA', 'P', 'LA'):
                            src_img = src_img.convert('RGB')
                        w, h = src_img.size
                        scale = min(1.0, 520 / w, 400 / h)
                        new_w = int(w * scale)
                        new_h = int(h * scale)
                        buf = io.BytesIO()
                        src_img.resize((new_w, new_h), PILImage.LANCZOS).save(
                            buf, format='JPEG', quality=90)
                    buf.seek(0)
                    xl_img        = XLImage(buf)
                    xl_img.width  = new_w
                    xl_img.height = new_h
                    ws.add_image(xl_img, anchor)
                    issue.was_corrected = True
                    modified = True
                except Exception as e:
                    print(f'[CORRECTOR] Rule 3 photo re-embed failed: {e}', flush=True)

            # ── Rule 4: write exact special value ────────────────────────────
            elif issue.rule == 4:
                ws = None
                if sheet_tag == 'Appendix-A':
                    ws = _find_ws(wb, 'Appendix-A')
                elif sheet_tag == 'Appendix-B':
                    ws = _find_ws(wb, 'Appendix-B')
                if ws is not None:
                    _write_cell(ws, addr, issue.expected)
                    issue.was_corrected = True
                    modified = True

            # ── Rule 5: write value extracted from raw input ─────────────────
            elif issue.rule == 5:
                ws = _find_ws(wb, 'Appendix-A')
                if ws is not None and issue.expected and not issue.expected.startswith('('):
                    _write_cell(ws, addr, issue.expected)
                    issue.was_corrected = True
                    modified = True

        except Exception as exc:
            print(f'[CORRECTOR] Failed to fix {issue.cell_or_field}: {exc}', flush=True)

    if modified:
        wb.save(report_path)
        print(f'[CORRECTOR] Saved corrections to {report_path}', flush=True)
    wb.close()


# ─────────────────────────────────────────────────────────────────────────────
#  correct_report() entry point
# ─────────────────────────────────────────────────────────────────────────────

def correct_report(report_path: str, report_json: dict, fmt: str,
                   result: CheckResult) -> str:
    """
    Fix every auto_fixable issue in `result`.

    After fixes are applied, runs check_report() once more on the corrected file
    and logs any residual issues (no second correction pass attempted).

    Returns the path of the corrected file (same path — overwritten in-place).
    """
    fixable = [i for i in result.issues if i.auto_fixable]
    if not fixable:
        return report_path

    try:
        if fmt in ('excel_rb', 'excel_amc'):
            _correct_excel(report_path, report_json, fmt, fixable)
        # Word auto-fix not implemented (all word issues have auto_fixable=False)
    except Exception as exc:
        print(f'[CORRECTOR] correct_report failed: {exc}', flush=True)
        return report_path

    # Post-correction check — log residual issues but do not re-correct
    try:
        post_result = check_report(report_path, report_json, fmt)
        for issue in post_result.issues:
            print(
                f'[CHECKER] Residual after correction: rule={issue.rule} '
                f'cell={issue.cell_or_field} found={issue.found!r}',
                flush=True,
            )
    except Exception:
        pass

    return report_path


# ─────────────────────────────────────────────────────────────────────────────
#  Logging
# ─────────────────────────────────────────────────────────────────────────────

def log_issues(result: CheckResult, phone: str, bridge_name: str, fmt: str) -> None:
    """Append all issues (and whether each was corrected) to checker_errors.log."""
    if not result.issues:
        return
    try:
        with open(LOG_PATH, 'a', encoding='utf-8') as f:
            for issue in result.issues:
                ts = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')
                corrected = 'YES' if issue.was_corrected else 'NO'
                f.write(
                    f"[{ts} UTC] phone={phone} bridge={bridge_name} fmt={fmt}\n"
                    f"  RULE {issue.rule} | {issue.severity.upper()} | "
                    f"cell={issue.cell_or_field} | found={issue.found!r} | "
                    f"expected={issue.expected!r} | auto_corrected={corrected}\n"
                )
    except Exception as exc:
        print(f'[CHECKER] log_issues failed: {exc}', flush=True)
